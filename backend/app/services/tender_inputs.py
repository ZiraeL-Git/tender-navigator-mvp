from __future__ import annotations

from io import BytesIO
from urllib.parse import parse_qs, urlparse

from docx import Document

from backend.app.api.schemas import TenderInputImportRequest
from backend.app.repositories.storage import StorageRepository
from backend.app.services.file_storage import FileStorageService


class TenderInputService:
    def __init__(
        self,
        storage: StorageRepository,
        file_storage: FileStorageService,
    ) -> None:
        self.storage = storage
        self.file_storage = file_storage

    async def create_manual_upload_input(
        self,
        *,
        company_profile_id: int,
        files,
    ) -> dict:
        folder_name = self.file_storage.create_folder_name("manual-upload")
        documents = await self.file_storage.save_uploads(folder_name, files)
        title = ", ".join(document["filename"] for document in documents) or "Manual upload"

        return self.storage.create_tender_input(
            payload={
                "company_profile_id": company_profile_id,
                "source_type": "manual_upload",
                "source_value": title,
                "source_url": None,
                "notice_number": None,
                "title": title,
                "customer_name": None,
                "deadline": None,
                "max_price": None,
                "status": "imported",
                "normalized_payload": {
                    "source_type": "manual_upload",
                    "document_names": [document["filename"] for document in documents],
                },
                "documents": documents,
                "last_error": None,
            }
        )

    def import_from_reference(self, payload: TenderInputImportRequest) -> dict:
        source_url = payload.source_url
        notice_number = payload.notice_number or self._extract_notice_number_from_url(source_url)
        source_type = "notice_number" if notice_number else "source_url"
        source_value = notice_number or source_url or "external-source"
        title = payload.title or (f"Закупка {notice_number}" if notice_number else "Импортированная закупка")

        normalized_payload = {
            "source_type": source_type,
            "source_value": source_value,
            "source_url": source_url,
            "notice_number": notice_number,
            "title": title,
            "customer_name": payload.customer_name,
            "deadline": payload.deadline,
            "max_price": payload.max_price,
        }

        generated_document = self.file_storage.save_generated_file(
            self.file_storage.create_folder_name("tender-input"),
            filename=(f"tender-card-{notice_number or 'source'}.docx"),
            content=self._build_card_docx(normalized_payload),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            kind="generated_card",
        )

        return self.storage.create_tender_input(
            payload={
                "company_profile_id": payload.company_profile_id,
                "source_type": source_type,
                "source_value": source_value,
                "source_url": source_url,
                "notice_number": notice_number,
                "title": title,
                "customer_name": payload.customer_name,
                "deadline": payload.deadline,
                "max_price": payload.max_price,
                "status": "imported",
                "normalized_payload": normalized_payload,
                "documents": [generated_document],
                "last_error": None,
            }
        )

    def _extract_notice_number_from_url(self, source_url: str | None) -> str | None:
        if not source_url:
            return None

        parsed = urlparse(source_url)
        params = parse_qs(parsed.query)
        for key in ("purchaseNumber", "noticeNumber", "number"):
            values = params.get(key)
            if values:
                return values[0]

        digits = "".join(char for char in source_url if char.isdigit())
        if 11 <= len(digits) <= 19:
            return digits
        return None

    def _build_card_docx(self, payload: dict) -> bytes:
        document = Document()
        document.add_heading("Карточка закупки", level=1)
        document.add_paragraph(f"Источник: {payload['source_type']}")
        document.add_paragraph(f"Значение источника: {payload['source_value']}")
        if payload.get("source_url"):
            document.add_paragraph(f"Ссылка: {payload['source_url']}")
        if payload.get("notice_number"):
            document.add_paragraph(f"Номер извещения: {payload['notice_number']}")
        document.add_paragraph(f"Наименование закупки: {payload['title']}")
        if payload.get("customer_name"):
            document.add_paragraph(f"Наименование организации: {payload['customer_name']}")
        if payload.get("deadline"):
            document.add_paragraph(f"Дата и время окончания срока подачи заявок: {payload['deadline']}")
        if payload.get("max_price"):
            document.add_paragraph(f"Начальная (максимальная) цена контракта: {payload['max_price']}")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
