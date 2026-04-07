from docx import Document
from pypdf import PdfReader

from schemas import DocumentType, TenderDocument
from services.text_utils import normalize_text


def detect_document_type(filename: str, text: str = "") -> DocumentType:
    name = filename.lower()
    head = normalize_text(text[:1500]).lower()

    if "извещ" in name or "извещение о проведении закупки" in head or "номер извещения" in head:
        return DocumentType.notice

    if "тз" in name or "техническ" in name or "техническое задание" in head:
        return DocumentType.spec

    if "контракт" in name or "договор" in name or "проект договора" in head:
        return DocumentType.contract

    if "прилож" in name:
        return DocumentType.attachment

    return DocumentType.other


def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)

    joined = "\n".join(parts).strip()

    if len(joined) < 50:
        return ""

    return joined


def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_text_from_uploaded_file(uploaded_file) -> str:
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    if filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)

    return ""


def build_tender_documents(uploaded_files) -> list[TenderDocument]:
    documents: list[TenderDocument] = []

    for uploaded_file in uploaded_files:
        text = extract_text_from_uploaded_file(uploaded_file)
        doc_type = detect_document_type(uploaded_file.name, text)

        documents.append(
            TenderDocument(
                filename=uploaded_file.name,
                doc_type=doc_type,
                extracted_text=text,
                text_length=len(text),
            )
        )

    return documents


def combine_documents_text(documents: list[TenderDocument]) -> str:
    parts = []

    for doc in documents:
        part = f"""
===== ДОКУМЕНТ =====
Имя файла: {doc.filename}
Тип документа: {doc.doc_type.value}

{doc.extracted_text or ""}
"""
        parts.append(part.strip())

    return "\n\n".join(parts)